import os
import uuid
from typing import List, Dict, Any, Optional
from pydantic import BaseModel, Field
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse

from typing_extensions import TypedDict
from langchain_ollama import ChatOllama
from langgraph.graph import StateGraph, START, END
from langgraph.types import interrupt, Command
from langgraph.checkpoint.memory import MemorySaver
from docx import Document


llm = ChatOllama(model="llama3.2")

# --- Industry-Standard Business Templates ---
DOCUMENT_TEMPLATES = {
    "sop": [
        "1. Document Control & Revision History",
        "2. Purpose & Objective",
        "3. Scope of Application",
        "4. Roles and Responsibilities",
        "5. Step-by-Step Standard Procedure",
        "6. Troubleshooting & Exceptions"
    ],
    "proposal": [
        "1. Executive Summary",
        "2. Problem Statement & Market Need",
        "3. Proposed Solution & Technical Approach",
        "4. Project Timeline & Milestones",
        "5. Operational Cost & Budget Breakdown",
        "6. Conclusion & Next Steps"
    ],
    "meeting_minutes": [
        "1. Meeting Metadata (Date, Time, Attendees)",
        "2. Meeting Objectives",
        "3. Core Discussion Topics & Debates",
        "4. Decisions Made & Action Items"
    ],
    "technical_design": [
        "1. Abstract & System Goals",
        "2. Architectural Overview (High-Level Design)",
        "3. Component Deep Dive & Data Models",
        "4. Performance, Scalability & Security Constraints"
    ]
}

# --- AGENT STATE ---
class AgentState(TypedDict):
    request: str
    document_type: str
    plan: List[str]
    current_step_index: int
    document_sections: Dict[str, str]
    latest_draft: str          # Tracks the unapproved section for review
    feedback: str              # Holds the user's feedback string for the section
    file_path: str
    logs: List[str]

# --- NODES ---

def planner_node(state: AgentState) -> Dict[str, Any]:
    logs = state.get("logs", []) + ["Planner activated."]
    request_lower = state["request"].lower()
    
    chosen_type = "custom"
    plan = []
    
    for doc_type, sections in DOCUMENT_TEMPLATES.items():
        if doc_type.replace("_", " ") in request_lower or doc_type in request_lower:
            chosen_type = doc_type
            plan = sections.copy()
            logs.append(f"Matched request to structured template: [{doc_type.upper()}]")
            break
            
    if chosen_type == "custom":
        logs.append("No template found. Routing to LLM for dynamic outline planning...")
        prompt = (
            f"Analyze this business document request: '{state['request']}'.\n"
            f"Generate a clean numbered outline of sections required to fulfill it.\n"
            f"Respond ONLY with a numbered list of section headers, one per line."
        )
        response = llm.invoke(prompt)
        plan = [line.strip() for line in response.content.split("\n") if line.strip() and (line.strip()[0].isdigit() or line.strip().startswith("-"))]
        
        if not plan:
            plan = ["1. Executive Summary", "2. Detailed Overview", "3. Conclusion"]
    
    return {
        "document_type": chosen_type,
        "plan": plan,
        "current_step_index": 0,
        "document_sections": {},
        "latest_draft": "",
        "feedback": "",
        "logs": logs + [f"Finalized structural outline with {len(plan)} sections."]
    }

def executor_node(state: AgentState) -> Dict[str, Any]:
    """Generates or rewrites content for the active section."""
    plan = state["plan"]
    idx = state["current_step_index"]
    current_section = plan[idx]
    
    sections_completed = list(state["document_sections"].keys())
    context_clause = ""
    if sections_completed:
        context_clause = f"You have already successfully written content for: {', '.join(sections_completed)}. Do not repeat points covered there."

    # --- FIX: Inject the previous draft alongside the feedback ---
    if state.get("feedback") and state.get("latest_draft"):
        context_clause += (
            f"\n\n--- REVISION MODE ---"
            f"\nYou are revising an existing draft based on human feedback."
            f"\nHere is your PREVIOUS DRAFT for this section:"
            f"\n\"\"\"\n{state['latest_draft']}\n\"\"\""
            f"\n\nCRITICAL USER FEEDBACK: '{state['feedback']}'"
            f"\n\nTask: Rewrite the previous draft, modifying it *only* to satisfy the user's feedback. Maintain the deterministic structure but execute the requested edits perfectly."
        )
        
    prompt = (
        f"You are a professional business writer drafting a {state['document_type'].upper()}.\n"
        f"Overall Objective: '{state['request']}'\n\n"
        f"Task: Write comprehensive, detailed body text for this section:\n"
        f"👉 '{current_section}'\n\n"
        f"{context_clause}\n"
        f"Requirements:\n"
        f"- Focus deeply on the scope of *this section only*.\n"
        f"- Do not include the section title header or conversational introductions."
    )
    
    response = llm.invoke(prompt)
    return {
        "latest_draft": response.content.strip(),
        "logs": state["logs"] + [f"Generated draft for section: {current_section}"]
    }

def human_review_node(state: AgentState) -> Dict[str, Any]:
    """Triggers a clean interrupt to pause the graph execution state and get user choice."""
    plan = state["plan"]
    idx = state["current_step_index"]
    current_section = plan[idx]

    # Clean interrupt baseline
    user_input = interrupt({
        "section_title": current_section,
        "drafted_content": state["latest_draft"],
        "message": f"Review section '{current_section}'. Provide revision feedback or send an empty string to approve."
    })
    
    # Receive input payload via Command(resume=...)
    received_feedback = user_input.get("feedback", "").strip()
    
    return {
        "feedback": received_feedback
    }

def evaluation_router_node(state: AgentState) -> Dict[str, Any]:
    """Evaluates the human interaction results and modifies indexing tracking coordinates."""
    plan = state["plan"]
    idx = state["current_step_index"]
    current_section = plan[idx]
    
    if not state["feedback"]:
        # User approved the section: Save draft to master registry and advance index counter
        updated_sections = dict(state["document_sections"])
        updated_sections[current_section] = state["latest_draft"]
        return {
            "document_sections": updated_sections,
            "current_step_index": idx + 1,
            "feedback": "",
            "latest_draft": "",
            "logs": state["logs"] + [f"Approved section {idx+1}/{len(plan)}: {current_section}"]
        }
    else:
        # User requested changes: Clear state variables but leave index alone to trigger a rebuild loop
        return {
            "logs": state["logs"] + [f"Feedback recorded for section '{current_section}': {state['feedback']}"]
        }

def compiler_node(state: AgentState) -> Dict[str, Any]:
    doc = Document()
    doc.add_heading(f"Professional Business Output: {state['document_type'].replace('_', ' ').title()}", level=0)
    
    for section_title, content in state["document_sections"].items():
        doc.add_heading(section_title, level=1)
        doc.add_paragraph(content)
        doc.add_paragraph("") 
        
    filename = f"compiled_agent_doc_{uuid.uuid4().hex[:8]}.docx"
    file_path = os.path.join(os.getcwd(), filename)
    print(f"document saved at: {file_path}")
    doc.save(file_path)
    
    return {
        "file_path": file_path,
        "logs": state["logs"] + [f"Document compiled successfully."]
    }

# --- CONTROL FLOW ROUTING ---
def route_after_evaluation(state: AgentState):
    """Determines whether to keep writing sections, loop back for feedback modifications, or finalize document."""
    # If the user just left a piece of feedback, loop directly back to executor to rewrite the current step
    if state.get("feedback"):
        return "loop_to_rewrite"
        
    # If no feedback was left, look to see if there are steps remaining in the execution plan array
    if state["current_step_index"] < len(state["plan"]):
        return "continue_to_next"
        
    return "finalize_document"

# --- RECONFIGURED WORKFLOW GRAPH ---
workflow = StateGraph(AgentState)
workflow.add_node("planner", planner_node)
workflow.add_node("executor", executor_node)
workflow.add_node("human_review", human_review_node)
workflow.add_node("evaluation_router", evaluation_router_node)
workflow.add_node("compiler", compiler_node)

workflow.add_edge(START, "planner")
workflow.add_edge("planner", "executor")
workflow.add_edge("executor", "human_review")
workflow.add_edge("human_review", "evaluation_router")

# Evaluate state routes using deterministic conditional mappings
workflow.add_conditional_edges(
    "evaluation_router",
    route_after_evaluation,
    {
        "loop_to_rewrite": "executor",
        "continue_to_next": "executor",
        "finalize_document": "compiler"
    }
)
workflow.add_edge("compiler", END)

memory = MemorySaver()
agent_graph = workflow.compile(checkpointer=memory)